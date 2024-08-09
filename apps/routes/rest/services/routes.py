#!/usr/bin/env python
""" Configuration Tool

The Configuration Tool is a software program produced for the European Space
Agency.

The purpose of this tool is to keep under configuration control the changes
in the Ground Segment components of the Copernicus Programme, in the
framework of the Coordination Desk Programme, managed by Telespazio S.p.A.

This program is free software: you can redistribute it and/or modify it under
the terms of the GNU General Public License as published by the Free Software
Foundation, either version 3 of the License, or (at your option) any later
version.

This program is distributed in the hope that it will be useful, but WITHOUT
ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or FITNESS
FOR A PARTICULAR PURPOSE. See the GNU General Public License for more details.
You should have received a copy of the GNU General Public License along with
this program. If not, see <http://www.gnu.org/licenses/>.
"""

__author__ = "Coordination Desk Development Team"
__contact__ = "coordination_desk@telespazio.com"
__copyright__ = "Copyright 2024, Telespazio S.p.A."
__license__ = "GPLv3"
__status__ = "Production"
__version__ = "1.0.0"

import json

import docx
import html2text
import pymongo
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls
from docx.shared import Pt, Inches
from flask import Response, send_file
from flask import request
from flask_login import login_required

import apps.utils.auth_utils as auth_utils
import apps.utils.db_utils as db_utils
from apps.models.nosql.Graph import Graph
from apps.models.sql import Scenario
from apps.routes.rest.services import blueprint
from apps.utils.word_document_generator import WordGenerator


@blueprint.route('/rest/api/services/<config_id>', methods=['GET'])
@login_required
def get_services(config_id):
    try:
        graph = Graph()
        scen_graph = graph.find({'id': config_id})
        scen_graph = scen_graph[0]

        # Add the external field in the model
        json_data = json.loads(scen_graph['graph'])
        for index, service in enumerate(json_data['services']):
            if 'external' not in service:
                service['external'] = False
        updated_graph_string = json.dumps(json_data)
        scen_graph['graph'] = updated_graph_string
        result = graph.update_one({'id': config_id}, scen_graph)

        return Response(json.dumps(scen_graph, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)
    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/services', methods=['POST'])
@login_required
def add_service():
    """
    :return:
    :rtype:
    """
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()
        scen_graph = graph.find({'id': body['config_id']})
        scen_graph = scen_graph[0]

        json_data = json.loads(scen_graph['graph'])

        if 'services' not in json_data:
            json_data['services'] = []

        json_data['services'].append({
            'id': db_utils.generate_uuid(),
            'type': body['type'],
            'provider': body['provider'],
            'external': body['external'],
            'satellite_units': body['satellite_units'],
            'interface_point': body['interface_point'],
            'cloud_provider': body['cloud_provider'],
            'rolling_period': body['rolling_period'],
            'operational_ipfs': body['operational_ipfs'],
            'references': body['references']
        })

        updated_graph_string = json.dumps(json_data)
        scen_graph['graph'] = updated_graph_string
        result = graph.update_one({'id': body['config_id']}, scen_graph)

        if result is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(scen_graph, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/services', methods=['PUT'])
@login_required
def update_service():
    """
    :return:
    :rtype:
    """
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()
        scen_graph = graph.find({'id': body['config_id']})
        scen_graph = scen_graph[0]

        json_data = json.loads(scen_graph['graph'])

        for index, service in enumerate(json_data['services']):
            if service['id'] == body['id']:
                service['type'] = body['type']
                service['provider'] = body['provider']
                service['external'] = body['external']
                service['satellite_units'] = body['satellite_units']
                service['interface_point'] = body['interface_point']
                service['cloud_provider'] = body['cloud_provider']
                service['rolling_period'] = body['rolling_period']
                service['operational_ipfs'] = body['operational_ipfs']
                service['references'] = body['references']

        updated_graph_string = json.dumps(json_data)
        scen_graph['graph'] = updated_graph_string
        result = graph.update_one({'id': body['config_id']}, scen_graph)

        if result is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(scen_graph, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/services', methods=['DELETE'])
@login_required
def delete_service():
    """
    :return:
    :rtype:
    """
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()
        scen_graph = graph.find({'id': body['config_id']})
        scen_graph = scen_graph[0]

        json_data = json.loads(scen_graph['graph'])

        for index, service in enumerate(json_data['services']):
            if service['id'] == body['service_id']:
                json_data['services'].remove(service)

        updated_graph_string = json.dumps(json_data)
        scen_graph['graph'] = updated_graph_string
        result = graph.update_one({'id': body['config_id']}, scen_graph)

        if result is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(result, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/services/interfaces', methods=['POST'])
@login_required
def add_service_interface():
    """
    :return:
    :rtype:
    """
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()
        scen_graph = graph.find({'id': body['config_id']})
        scen_graph = scen_graph[0]

        json_data = json.loads(scen_graph['graph'])

        if 'interfaces' not in json_data:
            json_data['interfaces'] = []

        json_data['interfaces'].append({
            'id': db_utils.generate_uuid(),
            'source_service_id': body['source_service_id'],
            'target_service_id': body['target_service_id'],
            'satellite_units': body['satellite_units'],
            'status': body['status']
        })

        updated_graph_string = json.dumps(json_data)
        scen_graph['graph'] = updated_graph_string
        result = graph.update_one({'id': body['config_id']}, scen_graph)

        if result is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(scen_graph, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/services/interfaces', methods=['PUT'])
@login_required
def update_service_interface():
    """
    :return:
    :rtype:
    """
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()
        scen_graph = graph.find({'id': body['config_id']})
        scen_graph = scen_graph[0]

        json_data = json.loads(scen_graph['graph'])

        for index, interface in enumerate(json_data['interfaces']):
            if interface['id'] == body['id']:
                interface['source_service_id'] = body['source_service_id']
                interface['target_service_id'] = body['target_service_id']
                interface['satellite_units'] = body['satellite_units']
                interface['status'] = body['status']

        updated_graph_string = json.dumps(json_data)
        scen_graph['graph'] = updated_graph_string
        result = graph.update_one({'id': body['config_id']}, scen_graph)

        if result is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(scen_graph, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/services/interfaces', methods=['DELETE'])
@login_required
def delete_service_interface():
    """
    :return:
    :rtype:
    """
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()
        scen_graph = graph.find({'id': body['config_id']})
        scen_graph = scen_graph[0]

        json_data = json.loads(scen_graph['graph'])

        for index, interface in enumerate(json_data['interfaces']):
            if interface['id'] == body['interface_id']:
                json_data['interfaces'].remove(interface)

        updated_graph_string = json.dumps(json_data)
        scen_graph['graph'] = updated_graph_string
        result = graph.update_one({'id': body['config_id']}, scen_graph)

        if result is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(result, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/services/commit/<config_id>', methods=['GET'])
@login_required
def commit_by_config_id(config_id):
    """
    :return:
    :rtype:
    """
    try:

        graph = Graph()
        ver_graphs = graph.history_find({'id': config_id},
                                        [('last_modify', pymongo.DESCENDING), ('tag', pymongo.ASCENDING),
                                         ('n_ver', pymongo.ASCENDING)])

        if ver_graphs is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        for i, version in enumerate(ver_graphs):
            version['last_modify'] = version['last_modify'].strftime("%d/%m/%Y, %H:%M:%S")

        return Response(json.dumps(ver_graphs, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/services/commit/<config_id>/<n_ver>', methods=['GET'])
@login_required
def commit_by_config_id_and_n_ver(config_id, n_ver):
    """
    :return:
    :rtype:
    """
    try:

        graph = Graph()
        n_ver = int(n_ver)
        ver_graphs = graph.history_find({'$and': [{'id': config_id}, {'n_ver': n_ver}]},
                                        [('last_modify', pymongo.DESCENDING), ('n_ver', pymongo.ASCENDING)])

        if ver_graphs is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)
        elif len(ver_graphs) == 0:
            return Response(json.dumps({'error': '404'}), mimetype="application/json", status=404)

        for i, version in enumerate(ver_graphs):
            version['last_modify'] = version['last_modify'].strftime("%d/%m/%Y, %H:%M:%S")

        return Response(json.dumps(ver_graphs, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/services/commit/<config_id>/<tag>', methods=['GET'])
@login_required
def commit_by_config_id_and_tag(config_id, tag):
    """
    :return:
    :rtype:
    """
    try:

        graph = Graph()
        ver_graphs = graph.history_find({'$and': [{'id': config_id}, {'tag': tag.upper()}]},
                                        [('last_modify', pymongo.DESCENDING), ('n_ver', pymongo.ASCENDING)])

        if ver_graphs is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)
        elif len(ver_graphs) == 0:
            return Response(json.dumps({'error': '404'}), mimetype="application/json", status=404)

        for i, version in enumerate(ver_graphs):
            version['last_modify'] = version['last_modify'].strftime("%d/%m/%Y, %H:%M:%S")

        return Response(json.dumps(ver_graphs, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/services/commit', methods=['POST'])
@login_required
def commit_interfaces_configuration():
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()

        versioned_obj = graph.versioning(body['idScenario'], body['tag'], body.get('comment', ''))

        if versioned_obj is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(versioned_obj, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/services/document/<config_id>', methods=['GET'])
@login_required
def download_services_document(config_id):
    """
    :param config_id:
    :return:
    :rtype:
    """

    # Instantiate the report generator
    word_doc_generator = WordGenerator('apps/docs/services/CSC_ESA_Operational_Configuration - template.docx')

    # Retrieve the configuration scenario
    scenario = Scenario.get_scenario(config_id)

    # Load the home graph
    graph = Graph()
    scen_graph = graph.find({'id': config_id})
    scen_graph = scen_graph[0]

    # Collect all nodes and connections from the selected configuration
    json_data = json.loads(scen_graph['graph'])
    services = json_data['services']
    interfaces = json_data['interfaces']
    dump_services_description(word_doc_generator, services)

    # Add a dedicated section with the interfaces matrix for each satellite
    satellites = ['S5P', 'S3B', 'S3A', 'S2B', 'S2A', 'S1A']
    for i, satellite in enumerate(satellites):
        dump_interfaces_matrix(word_doc_generator, satellite, services, interfaces)

    # Save and export the generated document
    path = word_doc_generator.save(scenario.name)
    return send_file(path, as_attachment=True)


def dump_services_description(word_doc_generator, services):

    # Retrieve the last paragraph
    prev_paragraph = word_doc_generator.get_paragraph("The components can be grouped as follows:")

    # Append an introduction paragraph
    par = word_doc_generator.add_paragraph_after(prev_paragraph)

    # Retrieve section and set page width
    word_doc_generator.set_section_width_height(2, 10692130, 7560310)

    # Append table
    table_name = 'Services list'
    table = word_doc_generator.add_table(table_name, len(services) + 1, 8, par)
    table.style = "Table Grid"
    table.allow_autofit = True
    header_row = ['Service Type', 'Service Provider', 'Satellite Unit(s)', 'Interface Point', 'Cloud Provider',
                  'Rolling Period [days]', 'Operational IPFs', 'References']
    for i, header in enumerate(header_row):
        word_doc_generator.add_text_to_cell_table(table_name, 0, i, header)
        cell = table.rows[0].cells[i]
        tblCell = cell._tc
        tblCellProperties = tblCell.get_or_add_tcPr()
        clShading = OxmlElement('w:shd')
        clShading.set(qn('w:fill'), "B3ECF1")
        tblCellProperties.append(clShading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for j, service in enumerate(services):
        word_doc_generator.add_text_to_cell_table(table_name, j + 1, 0, service['type'])
        word_doc_generator.add_text_to_cell_table(table_name, j + 1, 1, service['provider'])
        word_doc_generator.add_text_to_cell_table(table_name, j + 1, 2, service['satellite_units'])
        word_doc_generator.add_text_to_cell_table(table_name, j + 1, 3, service['interface_point'])
        word_doc_generator.add_text_to_cell_table(table_name, j + 1, 4, service['cloud_provider'])
        word_doc_generator.add_text_to_cell_table(table_name, j + 1, 5, service['rolling_period'])
        word_doc_generator.add_text_to_cell_table(table_name, j + 1, 6, service['operational_ipfs'])
        word_doc_generator.add_text_to_cell_table(table_name, j + 1, 7, service['references'])
        for k, col in enumerate(header_row):
            cell = table.rows[j + 1].cells[k]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return


def dump_interfaces_matrix(word_doc_generator, satellite, services, interfaces):

    # Retrieve the last paragraph
    prev_paragraph = word_doc_generator.get_paragraph("CSC INTERFACE MATRIX")

    # Retrieve section and set page width
    word_doc_generator.set_section_width_height(2, 10692130, 7560310)

    # Define locally the title name
    heading_name = {'S1A': 'Sentinel-1A', 'S2A': 'Sentinel-2A', 'S2B': 'Sentinel-2B',
                    'S3A': 'Sentinel-3A', 'S3B': 'Sentinel-3B', 'S5P': 'Sentinel-5P'}

    # Increment level, to create a new child section
    par = word_doc_generator.add_paragraph_after(prev_paragraph, heading_name[satellite] + ' Interfaces Matrix', 'Heading02')

    # Collect services relevant to the selected satellite
    filtered_services = [service for service in services if satellite in service['satellite_units'] or
                         satellite[:2] in [x.strip() for x in service['satellite_units'].split(',')]]

    # Sort services on the basis of the service type
    filtered_services.sort(key=lambda service: service['type'])

    # Initialize the interface matrix
    table_name = satellite + ' interface matrix'
    table = word_doc_generator.add_table(table_name, len(filtered_services) + 1, len(filtered_services) + 1, par)
    table.style = "Table Grid"
    table.allow_autofit = True

    # Set the font size and alignment in whole table
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Add from / to services labels
    for i, service in enumerate(filtered_services):

        # Horizontal row with labels
        word_doc_generator.add_text_to_cell_table(table_name, 0, i + 1,
                                                  service['type'][:3] + ' - ' + service['provider'])
        cell = table.rows[0].cells[i + 1]
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
            if service['external']:
                shading_elm = parse_xml(r'<w:shd {} w:fill="EE6B6E"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)

        # Vertical column with labels
        word_doc_generator.add_text_to_cell_table(table_name, i + 1, 0,
                                                  service['type'][:3] + ' - ' + service['provider'])
        cell = table.rows[i + 1].cells[0]
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        if service['external']:
            shading_elm = parse_xml(r'<w:shd {} w:fill="EE6B6E"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)

    # Populate interface matrix with values
    # Loop over the interfaces, and fill the cells corresponding to the row / col
    for i, service_row in enumerate(filtered_services):
        for j, service_col in enumerate(filtered_services):
            for iff in interfaces:
                if iff['source_service_id'] == service_row['id'] and iff['target_service_id'] == service_col['id']:
                    val = '1' if iff['status'] == 'Operational' else '2'
                    cell = table.rows[i + 1].cells[j + 1]
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                    word_doc_generator.add_text_to_cell_table(table_name, i + 1, j + 1, val)
                    shading_elm = parse_xml(r'<w:shd {} w:fill="ABF7B1"/>'.format(nsdecls('w'))) \
                        if iff['status'] == 'Operational' \
                        else parse_xml(r'<w:shd {} w:fill="00FFFF"/>'.format(nsdecls('w')))
                    table.rows[i + 1].cells[j + 1]._tc.get_or_add_tcPr().append(shading_elm)

    return
