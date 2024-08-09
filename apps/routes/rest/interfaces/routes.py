#!/usr/bin/env python
""" Configuration Tool

The Configuration Tool is a software program produced for the European Space
Agency.

The purpose of this tool is to keep under configuration control the changes
in the Ground Segment components of the Copernicus Programme, in the
framework of the Coordination Desk Programme, managed by Telespazio S.p.A.

This program is free software: you can redistribute it and/or modify it under
the terms of the GNU General Public License as published by the Free Software
Foundation, either version 3 of the License, or (at your option) any later
version.

This program is distributed in the hope that it will be useful, but WITHOUT
ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or FITNESS
FOR A PARTICULAR PURPOSE. See the GNU General Public License for more details.
You should have received a copy of the GNU General Public License along with
this program. If not, see <http://www.gnu.org/licenses/>.
"""

__author__ = "Coordination Desk Development Team"
__contact__ = "coordination_desk@telespazio.com"
__copyright__ = "Copyright 2024, Telespazio S.p.A."
__license__ = "GPLv3"
__status__ = "Production"
__version__ = "1.0.0"

import json
import os
import urllib
import re

import html2text
import pymongo
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from flask import Response, send_file
from flask import request
from flask_login import login_required

import apps.models.sql.Scenario as Scenario
import apps.utils.auth_utils as auth_utils
import apps.utils.db_utils as db_utils
from apps.models.nosql.Graph import Graph
from apps.routes.rest.interfaces import blueprint
from apps.utils.file_utils import safe_open_w
from apps.utils.word_document_generator import WordGenerator


@blueprint.route('/rest/api/interfaces/<config_id>', methods=['GET'])
@login_required
def get_interfaces_configuration(config_id):
    try:
        graph = Graph()
        scen_graph = graph.find({'id': config_id})
        scen_graph = scen_graph[0]
        return Response(json.dumps(scen_graph, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)
    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/interfaces', methods=['PUT'])
@login_required
def save_interfaces_configuration():
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)

        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()
        obj = graph.find({'id': body['id']})
        ''''''
        obj = obj[0]
        graph_obj = {}
        '''
        if not isinstance(obj, dict):
            graph_obj = {'graph': json.loads(obj['graph'])}
        elif not isinstance(obj['graph'], dict):
            graph_obj = {'graph': json.loads(obj['graph'])}
        else:
            graph_obj = {'graph': obj['graph']}
        '''
        if not isinstance(body, dict):
            graph_obj = {'id': body['id'], 'graph': json.loads(body['graph'])}
        elif not isinstance(body['graph'], dict):
            graph_obj = {'id': body['id'], 'graph': json.loads(body['graph'])}
        else:
            graph_obj = {'id': body['id'], 'graph': body['graph']}
        # graph_obj = {'graph': json.loads(obj['graph'])}
        # jsonMerged = {**body, **graph_obj}
        # body = jsonMerged
        # body.update(graph_obj)
        body = graph_obj
        ''''''
        if obj is not None:
            graph.delete_many({'id': body['id']})

        obj = graph.insert(body)

        if obj is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(obj, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/interfaces/entity', methods=['POST'])
@login_required
def add_entity():
    """
    :return:
    :rtype:
    """
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()
        scen_graph = graph.find({'id': body['idScenario']})
        scen_graph = scen_graph[0]

        json_data = json.loads(scen_graph['graph'])
        json_data['nodes'].append({
            'id': db_utils.generate_uuid(),
            'name': body['name'],
            'external': body['external'],
            'description': body['description'],
            'locked': False,
            'positionX': 10,
            'positionY': 10,
            'endpoints': []
        })

        updated_graph_string = json.dumps(json_data)
        scen_graph['graph'] = updated_graph_string
        result = graph.update_one({'id': body['idScenario']}, scen_graph)

        if result is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(scen_graph, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/interfaces/entity', methods=['PUT'])
@login_required
def update_entity():
    """
    :return:
    :rtype:
    """
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()
        scen_graph = graph.find({'id': body['idScenario']})
        scen_graph = scen_graph[0]

        json_data = json.loads(scen_graph['graph'])
        n_list = json_data['nodes']
        updated_name = body['name']
        updated_descr = body['description']
        updated_ext = body['external']
        index_to_modify = next((i for i, n in enumerate(n_list) if n.get('id') == body['idFragment']), None)

        if index_to_modify is not None:
            json_data['nodes'][index_to_modify]['name'] = updated_name
            json_data['nodes'][index_to_modify]['description'] = updated_descr
            json_data['nodes'][index_to_modify]['external'] = updated_ext

        updated_graph_string = json.dumps(json_data)
        scen_graph['graph'] = updated_graph_string
        result = graph.update_one({'id': body['idScenario']}, scen_graph)

        if result is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(scen_graph, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/interfaces/entity', methods=['DELETE'])
@login_required
def delete_entity():
    """
    :return:
    :rtype:
    """
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()
        nodes = body['graph']
        json_data = json.loads(nodes)
        n_list = json_data['nodes']
        node = body['removedEntityId']
        index_to_delete = next((i for i, n in enumerate(n_list) if n.get('id') == node), None)
        print(f"Nodes before deletion: {n_list}")

        if index_to_delete is not None:
            del json_data['nodes'][index_to_delete]
            del body['removedEntityId']

        updated_graph_string = json.dumps(json_data)
        body['graph'] = updated_graph_string

        obj = graph.update_one({'id': body['id']}, body)
        if obj is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(obj, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/interfaces/interface', methods=['POST'])
@login_required
def add_interface():
    """
    :return:
    :rtype:
    """
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()
        scen_graph = graph.find({'id': body['idScenario']})
        scen_graph = scen_graph[0]

        json_data = json.loads(scen_graph['graph'])
        source = next((n for i, n in enumerate(json_data['nodes']) if n.get('id') == body['source']), None)
        target = next((n for i, n in enumerate(json_data['nodes']) if n.get('id') == body['target']), None)
        if source is not None and target is not None:
            source_uuid = db_utils.generate_uuid()
            if not 'endpoints' in source:
                source['endpoints'] = []
            source['endpoints'].append({
                'id': source_uuid,
                'uuid': source_uuid,
                'type': 'source'
            })
            target_uuid = db_utils.generate_uuid()
            if not 'endpoints' in target:
                target['endpoints'] = []
            target['endpoints'].append({
                'id': target_uuid,
                'uuid': target_uuid,
                'type': 'target'
            })
            conn_id = db_utils.generate_uuid()
            json_data['connections'].append({
                'id': conn_id,
                'name': body['name'],
                'source_ep_id': source_uuid,
                'target_ep_id': target_uuid,
                'source_entity_name': source['name'],
                'target_entity_name': target['name'],
                'impacted_elements': body['elements'],
                'description': body['description'],
                'protocol': body['protocol'],
                'content': body['content'],
                'references': body['references'],
                'notes': body['notes']
            })

        updated_graph_string = json.dumps(json_data)
        scen_graph['graph'] = updated_graph_string
        result = graph.update_one({'id': body['idScenario']}, scen_graph)

        if result is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(scen_graph, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/interfaces/interface', methods=['PUT'])
@login_required
def update_interface():
    """
    :return:
    :rtype:
    """
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()
        scen_graph = graph.find({'id': body['idScenario']})
        scen_graph = scen_graph[0]

        json_data = json.loads(scen_graph['graph'])
        for i, conn in enumerate(json_data['connections']):
            if conn['id'] == body['id']:
                conn['name'] = body['name']
                conn['impacted_elements'] = body['elements']
                conn['description'] = body['description']
                conn['protocol'] = body['protocol']
                conn['content'] = body['content']
                conn['references'] = body['references']
                conn['notes'] = body['notes']

        updated_graph_string = json.dumps(json_data)
        scen_graph['graph'] = updated_graph_string
        result = graph.update_one({'id': body['idScenario']}, scen_graph)

        if result is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(scen_graph, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/interfaces/interface', methods=['DELETE'])
@login_required
def delete_interface():
    """
    :return:
    :rtype:
    """
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()
        scen_graph = graph.find({'id': body['idScenario']})
        scen_graph = scen_graph[0]

        json_data = json.loads(scen_graph['graph'])
        source_ep_id = None
        target_ep_id = None
        for i, conn in enumerate(json_data['connections']):
            if conn['id'] == body['idInterface']:
                source_ep_id = conn['source_ep_id']
                target_ep_id = conn['target_ep_id']
                json_data['connections'].remove(conn)
        for j, node in enumerate(json_data['nodes']):
            if 'endpoints' not in node:
                continue
            for k, ep in enumerate(node['endpoints']):
                if ep['id'] == source_ep_id or ep['id'] == target_ep_id:
                    node['endpoints'].remove(ep)

        updated_graph_string = json.dumps(json_data)
        scen_graph['graph'] = updated_graph_string
        result = graph.update_one({'id': body['idScenario']}, scen_graph)

        if result is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(scen_graph, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/interfaces/commit/<config_id>', methods=['GET'])
@login_required
def commit_by_config_id(config_id):
    """
    :return:
    :rtype:
    """
    try:

        graph = Graph()
        ver_graphs = graph.history_find({'id': config_id},
                                        [('last_modify', pymongo.DESCENDING), ('tag', pymongo.ASCENDING),
                                         ('n_ver', pymongo.ASCENDING)])

        if ver_graphs is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)
        elif len(ver_graphs) == 0:
            return Response(json.dumps({'error': '404'}), mimetype="application/json", status=404)

        for i, version in enumerate(ver_graphs):
            version['last_modify'] = version['last_modify'].strftime("%d/%m/%Y, %H:%M:%S")

        return Response(json.dumps(ver_graphs, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/interfaces/commit/<config_id>/<n_ver>', methods=['GET'])
@login_required
def commit_by_config_id_and_n_ver(config_id, n_ver):
    """
    :return:
    :rtype:
    """
    try:

        graph = Graph()
        n_ver = int(n_ver)
        ver_graphs = graph.history_find({'$and': [{'id': config_id}, {'n_ver': n_ver}]},
                                        [('last_modify', pymongo.DESCENDING), ('n_ver', pymongo.ASCENDING)])

        if ver_graphs is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)
        elif len(ver_graphs) == 0:
            return Response(json.dumps({'error': '404'}), mimetype="application/json", status=404)

        for i, version in enumerate(ver_graphs):
            version['last_modify'] = version['last_modify'].strftime("%d/%m/%Y, %H:%M:%S")

        return Response(json.dumps(ver_graphs, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/interfaces/commit/<config_id>/<tag>', methods=['GET'])
@login_required
def commit_by_config_id_and_tag(config_id, tag):
    """
    :return:
    :rtype:
    """
    try:

        graph = Graph()
        ver_graphs = graph.history_find({'$and': [{'id': config_id}, {'tag': tag.upper()}]},
                                        [('last_modify', pymongo.DESCENDING), ('n_ver', pymongo.ASCENDING)])

        if ver_graphs is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)
        elif len(ver_graphs) == 0:
            return Response(json.dumps({'error': '404'}), mimetype="application/json", status=404)

        for i, version in enumerate(ver_graphs):
            version['last_modify'] = version['last_modify'].strftime("%d/%m/%Y, %H:%M:%S")

        return Response(json.dumps(ver_graphs, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/interfaces/commit', methods=['POST'])
@login_required
def commit_interfaces_configuration():
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)
        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        if body is None or len(body) == 0:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        graph = Graph()

        versioned_obj = graph.versioning(body['idScenario'], body['tag'], body.get('comment', ''))

        if versioned_obj is None:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        return Response(json.dumps(versioned_obj, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/interfaces/document/upload/<config_id>/<image_id>', methods=['POST'])
@login_required
def upload_image(config_id, image_id):
    try:
        if not auth_utils.is_user_authorized(['admin']):
            return Response(json.dumps("Not authorized", cls=db_utils.AlchemyEncoder), mimetype="application/json",
                            status=401)

        body = None
        if request.data != b'':
            body = json.loads(request.data.decode('utf-8'))
        else:
            return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)

        # Dump image to file
        # /apps/docs/interfaces/<config_id>/<entity_id>.png
        data_url = body['data_url']
        try:
            response = urllib.request.urlopen(data_url)
            with safe_open_w('apps/docs/interfaces/' + config_id + '/' + image_id + '.png', 'wb') as fd:
                fd.write(response.file.read())
                fd.close()
        except Exception as ex:
            print(ex)

        return Response(json.dumps(body, cls=db_utils.AlchemyEncoder), mimetype="application/json", status=200)

    except Exception as ex:
        return Response(json.dumps({'error': '500'}), mimetype="application/json", status=500)


@blueprint.route('/rest/api/interfaces/document/<config_id>', methods=['GET'])
@login_required
def download_interfaces_document(config_id):
    """
    :param config_id:
    :return:
    :rtype:
    """

    # Instantiate the report generator
    word_doc_generator = WordGenerator(
        'apps/docs/interfaces/[ESA-EOPG-EOPGC-IF-6] ESA EO Operations Framework (EOF) - CSC - Ground Segment Master '
        'ICD - template.docx')

    # Retrieve the configuration scenario
    scenario = Scenario.get_scenario(config_id)

    # Load the home graph
    graph = Graph()
    scen_graph = graph.find({'id': config_id})
    scen_graph = scen_graph[0]

    # Collect all nodes and connections from the selected configuration
    json_data = json.loads(scen_graph['graph'])
    nodes = json_data['nodes']
    connections = json_data['connections']

    # Initialize the prev_paragraph anchor, i.e. the position from where paragraph should be appended
    prev_paragraph = word_doc_generator.get_paragraph("CURRENT COPERNICUS MISSIONS")

    # For every node, dump a dedicated section, with the description of all connected interfaces
    for i, node in enumerate(nodes):

        # Skip external entities
        if node['external']:
            continue

        # Given the current node, retrieve the corresponding image and select the relevant interfaces
        image_path = 'apps/docs/interfaces/' + config_id + '/' + node['id'] + '.png'
        selected_connections = select_connections(node, connections)
        dump_entity_description(word_doc_generator, prev_paragraph, node, image_path, selected_connections)

    # Save and export the generated document
    path = word_doc_generator.save(scenario.name)
    return send_file(path, as_attachment=True)


def select_connections(node, connections):
    selected_connections = []
    if 'endpoints' not in node:
        return selected_connections
    for ep in node['endpoints']:
        for conn in connections:
            if ep['id'] == conn['source_ep_id'] or ep['id'] == conn['target_ep_id']:
                if conn not in selected_connections:
                    selected_connections.append(conn)
    return selected_connections


def dump_entity_description(word_doc_generator, prev_paragraph, node, image_path, selected_connections):

    # Create the new paragraph corresponding to the provided node
    par = word_doc_generator.add_paragraph_after(prev_paragraph, node['name'], 'Heading03')

    # Add the interface table
    if selected_connections:
        table_name = node['name'] + ' interfaces'
        table = word_doc_generator.add_table(table_name, len(selected_connections) + 1, 9, par)
        table.style = "Table Grid"
        table.allow_autofit = True
        header_row = ['Interface Name', 'Satellite Unit(s)', 'Description', 'Service Source', 'Service Destination',
                      'Content', 'Protocol', 'References', 'Notes']
        for i, header in enumerate(header_row):
            word_doc_generator.add_text_to_cell_table(table_name, 0, i, header)
            cell = table.rows[0].cells[i]
            tblCell = cell._tc
            tblCellProperties = tblCell.get_or_add_tcPr()
            clShading = OxmlElement('w:shd')
            clShading.set(qn('w:fill'), "B3ECF1")
            tblCellProperties.append(clShading)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        for j, conn in enumerate(selected_connections):
            word_doc_generator.add_text_to_cell_table(table_name, j + 1, 0, conn['name'])
            word_doc_generator.add_text_to_cell_table(table_name, j + 1, 1, conn['impacted_elements'])
            word_doc_generator.add_text_to_cell_table(table_name, j + 1, 2, conn['description'])
            word_doc_generator.add_text_to_cell_table(table_name, j + 1, 3, conn['source_entity_name'])
            word_doc_generator.add_text_to_cell_table(table_name, j + 1, 4, conn['target_entity_name'])
            word_doc_generator.add_text_to_cell_table(table_name, j + 1, 5, conn['content'])
            word_doc_generator.add_text_to_cell_table(table_name, j + 1, 6, conn['protocol'])
            word_doc_generator.add_text_to_cell_table(table_name, j + 1, 7, conn['references'])
            word_doc_generator.add_text_to_cell_table(table_name, j + 1, 8, conn['notes'])
            for k, col in enumerate(header_row):
                cell = table.rows[j + 1].cells[k]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # Add the node picture
    if os.path.isfile(image_path):
        word_doc_generator.add_picture(image_path, par, node['name'])

    return
