#!/usr/bin/env python
""" Configuration Tool

The Configuration Tool is a software program produced for the European Space
Agency.

The purpose of this tool is to keep under configuration control the changes
in the Ground Segment components of the Copernicus Programme, in the
framework of the Coordination Desk Programme, managed by Telespazio S.p.A.

This program is free software: you can redistribute it and/or modify it under
the terms of the GNU General Public License as published by the Free Software
Foundation, either version 3 of the License, or (at your option) any later
version.

This program is distributed in the hope that it will be useful, but WITHOUT
ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or FITNESS
FOR A PARTICULAR PURPOSE. See the GNU General Public License for more details.
You should have received a copy of the GNU General Public License along with
this program. If not, see <http://www.gnu.org/licenses/>.
"""

__author__ = "Coordination Desk Development Team"
__contact__ = "coordination_desk@telespazio.com"
__copyright__ = "Copyright 2024, Telespazio S.p.A."
__license__ = "GPLv3"
__status__ = "Production"
__version__ = "1.0.0"

import docx
from htmldocx import HtmlToDocx
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

import tempfile

from docx.text.paragraph import Paragraph

from apps.utils import file_utils


class WordGenerator:

    def __init__(self, path_to_document=None):
        self.__document = None
        self.__map = {}
        self.create(path_to_document)
        return

    def create(self, path_to_document=None):
        """
        :return:
        :rtype:
        """
        if path_to_document:
            self.__document = Document(path_to_document)
        else:
            self.__document = Document()
        return

    def add_title_header(self, title):
        """
        :param title:
        :type title:
        :return:
        :rtype:
        """
        section = self.__document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = title
        return header

    def add_heading(self, name, level):
        """
        :param name:
        :type name:
        :param level:
        :type level:
        :return:
        :rtype:
        """
        self.__map[name] = self.__document.add_heading(name, level)
        return self.__map[name]

    def add_paragraph(self, name, text):
        """
        :param name:
        :type name:
        :param text:
        :type text:
        :return:
        :rtype:
        """
        self.__map[name] = self.__document.add_paragraph(text)
        return self.__map[name]

    def add_paragraph_before(self, paragraph, text=None, style=None):
        for index, par in enumerate(self.__document.paragraphs):
            if par.text.lower() == paragraph.text.lower():
                added_par = par.insert_paragraph_before(text, style)
        return None

    def add_paragraph_after(self, paragraph, text=None, style=None):
        for index, par in enumerate(self.__document.paragraphs):
            if par.text.lower() == paragraph.text.lower():
                next_par = self.__document.paragraphs[index + 1]
                added_par = next_par.insert_paragraph_before(text, style)
                return added_par
        return None

    def add_picture(self, image_path, paragraph=None, caption=None):
        """
        :param paragraph:
        :param image_data:
        :type image_data:
        :return:
        :rtype:
        """
        # skip function if image is null
        if image_path is None:
            return

        # adjust the size as needed
        width, height = 5, 4.5

        # Use a unique identifier for each image
        image_id = len(self.__map) + 1
        if paragraph is None:
            self.__map[image_id] = self.__document.add_picture(image_path, width=Inches(width),
                                                               height=Inches(height))
            last_paragraph = self.__document.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            added_par = Paragraph(new_p, paragraph._parent)
            run = added_par.add_run()
            self.__map[image_id] = run.add_picture(image_path, width=Inches(width),
                                                   height=Inches(height))
            added_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # If present, add a figure caption
            if caption:

                # New caption paragraph
                new_p = OxmlElement("w:p")
                added_par._p.addnext(new_p)
                caption_par = Paragraph(new_p, added_par._parent)
                caption_par.style = 'Caption'
                caption_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Add figure label
                caption_par.add_run('Figure ')

                # Add automatic numbering
                run = caption_par.add_run()
                r = run._r

                fldChar = docx.oxml.OxmlElement('w:fldChar')
                fldChar.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
                r.append(fldChar)

                instrText = docx.oxml.OxmlElement('w:instrText')
                instrText.text = ' SEQ Figure \\* ARABIC'
                r.append(instrText)

                fldChar = docx.oxml.OxmlElement('w:fldChar')
                fldChar.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
                r.append(fldChar)

                # Add text
                caption_par.add_run(' ' + caption)

        return image_id

    def add_table(self, name, rows, cols, paragraph=None):
        """
        :param paragraph:
        :param name:
        :type name:
        :param rows:
        :type rows:
        :param cols:
        :type cols:
        :return:
        :rtype:
        """
        self.__map[name] = self.__document.add_table(rows=rows, cols=cols)
        if paragraph is not None:
            paragraph._p.addnext(self.__map[name]._tbl)
        return self.__map[name]

    def add_text_to_cell_table(self, table_name, row, col, text):
        """
        :param table_name:
        :type table_name:
        :param row:
        :type row:
        :param col:
        :type col:
        :param text:
        :type text:
        :return:
        :rtype:
        """
        cell = self.__map[table_name].cell(row, col)
        parser = HtmlToDocx()
        parser.add_html_to_cell(text, cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
        return cell

    def set_horizontal_layout(self):
        sections = self.__document.sections
        for section in sections:
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height

    def get_paragraph(self, paragraph_name):
        for index, par in enumerate(self.__document.paragraphs):
            if par.text.lower() == paragraph_name.lower():
                return par
        return None

    def set_section_width_height(self, index, width, height):
        section = self.__document.sections[index]
        section.page_width = width
        section.page_height = height

    def save(self, name):
        """
        :param name:
        :type name:
        :return:
        :rtype:
        """
        path = tempfile.gettempdir()
        from apps.utils import auth_utils as utils
        self.__document.save(path + '/' + name + ' - ' + file_utils.get_date_for_file() + '.docx')
        return path + '/' + name + ' - ' + file_utils.get_date_for_file() + '.docx'
